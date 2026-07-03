#!/usr/bin/env python3
"""pandoc が出力した docx の表に罫線・ヘッダ装飾を加え、本文/余白も調整して読みやすくする。

usage:
    python style_tables.py <input.docx> <output.docx>

pandoc 既定の docx は表に罫線がなく、A4 縦の既定余白も広いため、
日本語 + 多列テーブルが極端に読みづらい。本スクリプトで以下を補正する:

- 全テーブルの全セルに単線（黒・幅 0.5pt）罫線を追加
- 1 行目（ヘッダ）に薄グレー (#DDDDDD) 背景 + 太字を適用
- セル内余白を上下 40 twips / 左右 80 twips に詰める
- テーブル全体を中央寄せ
- 本文段落フォントを 10pt、テーブル内フォントを 9pt に圧縮
- A4 縦のまま左右余白 15mm / 上下余白 18mm に縮小（既定 25mm から削減し横幅 +20mm 確保）

依存: `python-docx` (uv venv 等にインストールして実行)
"""

import sys

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

SRC = sys.argv[1]
DST = sys.argv[2]


def _set_cell_border(cell, color="000000", sz="4"):
    """セル四辺に単線罫線を設定。sz は 1/8 pt 単位 (4 → 0.5pt)。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        existing = tcBorders.find(tag)
        if existing is not None:
            tcBorders.remove(existing)
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)


def _set_cell_shading(cell, fill_hex):
    """セル背景色 (RGB hex without #)。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_margin(cell, top=40, left=80, bottom=40, right=80):
    """セル内余白 (twips = 1/20 pt)。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        tag = qn(f"w:{side}")
        existing = tcMar.find(tag)
        if existing is not None:
            tcMar.remove(existing)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)


def style_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                _set_cell_border(cell)
                _set_cell_margin(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if is_header:
                    _set_cell_shading(cell, "DDDDDD")
                for p in cell.paragraphs:
                    for run in p.runs:
                        if is_header:
                            run.font.bold = True
                        if run.font.size is None or run.font.size > Pt(10):
                            run.font.size = Pt(9)


def shrink_margins(doc):
    """A4 縦のまま余白を圧縮して横幅を確保。"""
    for section in doc.sections:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)


def shrink_normal_paragraph_size(doc):
    """本文段落のフォントサイズを 10pt に。見出しはスキップ。"""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") or p.style.name in ("Title", "Subtitle"):
            continue
        for run in p.runs:
            if run.font.size is None or run.font.size > Pt(10):
                run.font.size = Pt(10)


def main():
    doc = Document(SRC)
    shrink_margins(doc)
    shrink_normal_paragraph_size(doc)
    style_tables(doc)
    doc.save(DST)
    print(f"saved: {DST}")
    print(f"tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
